"""
DOCX Parser Module
=================

Optimized DOCX parsing with text extraction and embedded image handling.
"""

from typing import Optional, Dict, Any, List, Tuple
from pathlib import Path
from dataclasses import dataclass
import io
import zipfile

try:
    from docx import Document
    from docx.document import Document as DocumentType
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import _Cell, Table
    from docx.text.paragraph import Paragraph
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

from loguru import logger


@dataclass
class DOCXMetadata:
    """Metadata extracted from DOCX."""
    title: Optional[str] = None
    author: Optional[str] = None
    subject: Optional[str] = None
    keywords: Optional[str] = None
    created: Optional[str] = None
    modified: Optional[str] = None
    last_modified_by: Optional[str] = None
    revision: Optional[str] = None
    total_paragraphs: int = 0
    total_tables: int = 0
    total_images: int = 0
    has_embedded_images: bool = False


@dataclass
class EmbeddedImage:
    """Information about an embedded image."""
    name: str
    format: str
    size_bytes: int
    width: Optional[int] = None
    height: Optional[int] = None
    data: Optional[bytes] = None


@dataclass
class DOCXParseResult:
    """Result of DOCX parsing operation."""
    text_content: str
    metadata: DOCXMetadata
    paragraph_texts: List[str]
    table_texts: List[str]
    embedded_images: List[EmbeddedImage]
    structured_content: Dict[str, Any]
    confidence_score: float = 0.0
    extraction_method: str = "python_docx"
    error_message: Optional[str] = None


class DOCXParser:
    """
    Optimized DOCX parser with comprehensive content extraction.
    """
    
    def __init__(self):
        """Initialize DOCX parser."""
        if not PYTHON_DOCX_AVAILABLE:
            raise ImportError("python-docx is not available")
        
        self.pil_available = PIL_AVAILABLE
        logger.info(f"DOCX Parser initialized - PIL available: {self.pil_available}")
    
    def parse_docx(self, file_path: Path) -> DOCXParseResult:
        """
        Parse DOCX file and extract all content.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            DOCXParseResult with extracted content and metadata
        """
        logger.info(f"Starting DOCX parsing for {file_path}")
        
        try:
            doc = Document(file_path)
            
            # Extract metadata
            metadata = self._extract_metadata(doc, file_path)
            
            # Extract text content
            paragraph_texts = self._extract_paragraphs(doc)
            table_texts = self._extract_tables(doc)
            
            # Extract embedded images
            embedded_images = self._extract_embedded_images(file_path)
            metadata.total_images = len(embedded_images)
            metadata.has_embedded_images = len(embedded_images) > 0
            
            # Combine all text
            all_text = self._combine_text_content(paragraph_texts, table_texts)
            
            # Create structured content
            structured_content = self._create_structured_content(doc, paragraph_texts, table_texts)
            
            # Calculate confidence
            confidence = self._calculate_confidence(all_text, metadata)
            
            logger.info(f"DOCX parsing completed - {len(all_text)} characters, {len(embedded_images)} images")
            
            return DOCXParseResult(
                text_content=all_text,
                metadata=metadata,
                paragraph_texts=paragraph_texts,
                table_texts=table_texts,
                embedded_images=embedded_images,
                structured_content=structured_content,
                confidence_score=confidence,
                extraction_method="python_docx"
            )
            
        except Exception as e:
            logger.error(f"DOCX parsing failed: {e}")
            return DOCXParseResult(
                text_content="",
                metadata=DOCXMetadata(),
                paragraph_texts=[],
                table_texts=[],
                embedded_images=[],
                structured_content={},
                error_message=str(e)
            )
    
    def _extract_metadata(self, doc: Document, file_path: Path) -> DOCXMetadata:
        """Extract metadata from DOCX document."""
        try:
            core_props = doc.core_properties
            
            # Count document elements
            paragraph_count = len(doc.paragraphs)
            table_count = len(doc.tables)
            
            return DOCXMetadata(
                title=core_props.title,
                author=core_props.author,
                subject=core_props.subject,
                keywords=core_props.keywords,
                created=core_props.created.isoformat() if core_props.created else None,
                modified=core_props.modified.isoformat() if core_props.modified else None,
                last_modified_by=core_props.last_modified_by,
                revision=str(core_props.revision) if core_props.revision else None,
                total_paragraphs=paragraph_count,
                total_tables=table_count
            )
            
        except Exception as e:
            logger.warning(f"Failed to extract DOCX metadata: {e}")
            return DOCXMetadata()
    
    def _extract_paragraphs(self, doc: DocumentType) -> List[str]:
        """Extract text from all paragraphs."""
        paragraph_texts = []
        
        for paragraph in doc.paragraphs:
            try:
                text = paragraph.text.strip()
                if text:  # Only add non-empty paragraphs
                    paragraph_texts.append(text)
            except Exception as e:
                logger.warning(f"Failed to extract paragraph text: {e}")
                continue
        
        logger.debug(f"Extracted {len(paragraph_texts)} paragraphs")
        return paragraph_texts
    
    def _extract_tables(self, doc: DocumentType) -> List[str]:
        """Extract text from all tables."""
        table_texts = []
        
        for table_idx, table in enumerate(doc.tables):
            try:
                table_text = self._extract_table_text(table)
                if table_text.strip():
                    table_texts.append(table_text)
            except Exception as e:
                logger.warning(f"Failed to extract table {table_idx} text: {e}")
                continue
        
        logger.debug(f"Extracted {len(table_texts)} tables")
        return table_texts
    
    def _extract_table_text(self, table: Table) -> str:
        """Extract text from a single table."""
        table_data = []
        
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_data.append(cell_text)
            
            if any(row_data):  # Only add rows with content
                table_data.append(" | ".join(row_data))
        
        return "\n".join(table_data)
    
    def _extract_embedded_images(self, file_path: Path) -> List[EmbeddedImage]:
        """Extract embedded images from DOCX file."""
        images = []
        
        try:
            # DOCX files are ZIP archives
            with zipfile.ZipFile(file_path, 'r') as docx_zip:
                # Look for images in the media folder
                media_files = [f for f in docx_zip.namelist() if f.startswith('word/media/')]
                
                for media_file in media_files:
                    try:
                        image_data = docx_zip.read(media_file)
                        image_name = Path(media_file).name
                        image_format = Path(media_file).suffix.lower().lstrip('.')
                        
                        # Create image object
                        embedded_image = EmbeddedImage(
                            name=image_name,
                            format=image_format,
                            size_bytes=len(image_data),
                            data=image_data
                        )
                        
                        # Try to get image dimensions if PIL is available
                        if self.pil_available:
                            try:
                                with Image.open(io.BytesIO(image_data)) as img:
                                    embedded_image.width, embedded_image.height = img.size
                            except Exception as e:
                                logger.debug(f"Could not get dimensions for {image_name}: {e}")
                        
                        images.append(embedded_image)
                        
                    except Exception as e:
                        logger.warning(f"Failed to process embedded image {media_file}: {e}")
                        continue
        
        except Exception as e:
            logger.warning(f"Failed to extract embedded images: {e}")
        
        logger.info(f"Extracted {len(images)} embedded images")
        return images
    
    def _combine_text_content(self, paragraph_texts: List[str], table_texts: List[str]) -> str:
        """Combine all text content into a single string."""
        all_content = []
        
        # Add paragraph text
        if paragraph_texts:
            all_content.extend(paragraph_texts)
        
        # Add table text with separator
        if table_texts:
            all_content.append("\n--- TABLES ---\n")
            all_content.extend(table_texts)
        
        return "\n\n".join(all_content)
    
    def _create_structured_content(self, doc: DocumentType, 
                                 paragraph_texts: List[str], 
                                 table_texts: List[str]) -> Dict[str, Any]:
        """Create structured representation of document content."""
        structured = {
            "document_structure": [],
            "headings": [],
            "lists": [],
            "hyperlinks": []
        }
        
        try:
            # Process document elements in order
            for element in doc.element.body:
                if isinstance(element, CT_P):
                    # Paragraph
                    paragraph = Paragraph(element, doc)
                    text = paragraph.text.strip()
                    
                    if text:
                        # Check if it's a heading
                        if paragraph.style.name.startswith('Heading'):
                            structured["headings"].append({
                                "level": paragraph.style.name,
                                "text": text
                            })
                        
                        structured["document_structure"].append({
                            "type": "paragraph",
                            "style": paragraph.style.name,
                            "text": text
                        })
                
                elif isinstance(element, CT_Tbl):
                    # Table
                    structured["document_structure"].append({
                        "type": "table",
                        "text": "[TABLE CONTENT]"
                    })
            
            # Extract hyperlinks
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if run.element.xpath('.//w:hyperlink'):
                        structured["hyperlinks"].append(run.text)
        
        except Exception as e:
            logger.warning(f"Failed to create structured content: {e}")
        
        return structured
    
    def _calculate_confidence(self, text: str, metadata: DOCXMetadata) -> float:
        """Calculate confidence score for extraction."""
        confidence_factors = []
        
        # Factor 1: Text length and quality
        text_length = len(text.strip())
        if text_length > 1000:
            confidence_factors.append(0.9)
        elif text_length > 500:
            confidence_factors.append(0.8)
        elif text_length > 100:
            confidence_factors.append(0.6)
        elif text_length > 50:
            confidence_factors.append(0.4)
        else:
            confidence_factors.append(0.2)
        
        # Factor 2: Check for mostly special characters or whitespace
        if text_length > 0:
            alphanumeric_chars = sum(1 for c in text if c.isalnum())
            alphanumeric_ratio = alphanumeric_chars / text_length
            
            if alphanumeric_ratio > 0.7:
                confidence_factors.append(0.8)
            elif alphanumeric_ratio > 0.4:
                confidence_factors.append(0.6)
            else:
                confidence_factors.append(0.3)
        else:
            confidence_factors.append(0.0)
        
        # Factor 3: Document structure
        if metadata.total_paragraphs > 5:
            confidence_factors.append(0.8)
        elif metadata.total_paragraphs > 0:
            confidence_factors.append(0.6)
        else:
            confidence_factors.append(0.2)
        
        # Factor 3: Metadata presence
        metadata_score = 0
        if metadata.title:
            metadata_score += 0.2
        if metadata.author:
            metadata_score += 0.2
        if metadata.created:
            metadata_score += 0.1
        
        confidence_factors.append(min(metadata_score, 0.5))
        
        return sum(confidence_factors) / len(confidence_factors)
    
    def needs_ocr_fallback(self, text: str, embedded_images: List[EmbeddedImage]) -> bool:
        """
        Determine if DOCX needs OCR fallback processing.
        
        Args:
            text: Extracted text content
            embedded_images: List of embedded images
            
        Returns:
            True if OCR fallback is recommended
        """
        text_length = len(text.strip())
        
        # Check for insufficient text content
        if text_length < 50:
            return True
        
        # Check for mostly special characters
        if text_length > 0:
            alphanumeric_chars = sum(1 for c in text if c.isalnum())
            alphanumeric_ratio = alphanumeric_chars / text_length
            
            if alphanumeric_ratio < 0.3:  # Mostly special chars/whitespace
                return True
        
        # If document has many images but little text, might need OCR
        if len(embedded_images) > 3 and text_length < 200:
            return True
        
        # If we have images and very little text per image, might need OCR
        if len(embedded_images) > 0 and text_length / len(embedded_images) < 50:
            return True
        
        return False
    



def parse_docx(file_path: Path) -> DOCXParseResult:
    """
    Convenience function to parse a DOCX file.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        DOCXParseResult with extracted content
    """
    parser = DOCXParser()
    return parser.parse_docx(file_path)
