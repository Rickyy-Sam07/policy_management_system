#!/usr/bin/env python3
"""
Multi-Format Document Processor
Supports PDF, DOCX, and Email documents
"""

import os
import re
import email
import mimetypes
from typing import Dict, List, Any, Optional
from docx import Document
import fitz  # PyMuPDF
import requests

class MultiFormatProcessor:
    """Process PDF, DOCX, and Email documents"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc', '.eml', '.msg']
    
    def process_document_url(self, url: str) -> Dict[str, Any]:
        """Process document from URL - auto-detect format"""
        
        try:
            # Download document
            response = requests.get(url, timeout=30)
            response.raise_for_status()
            
            # Detect format from content-type or URL
            content_type = response.headers.get('content-type', '')
            file_ext = self._detect_format(url, content_type)
            
            if file_ext == '.pdf':
                return self._process_pdf_content(response.content)
            elif file_ext in ['.docx', '.doc']:
                return self._process_docx_content(response.content)
            elif file_ext in ['.eml', '.msg']:
                return self._process_email_content(response.content)
            else:
                return {'error': f'Unsupported format: {file_ext}'}
                
        except Exception as e:
            return {'error': f'Document processing failed: {str(e)}'}
    
    def _detect_format(self, url: str, content_type: str) -> str:
        """Detect document format"""
        
        # Check URL extension
        for ext in self.supported_formats:
            if url.lower().endswith(ext):
                return ext
        
        # Check content type
        if 'pdf' in content_type:
            return '.pdf'
        elif 'word' in content_type or 'officedocument' in content_type:
            return '.docx'
        elif 'email' in content_type or 'message' in content_type:
            return '.eml'
        
        return '.pdf'  # Default fallback
    
    def _process_pdf_content(self, content: bytes) -> Dict[str, Any]:
        """Process PDF content"""
        
        try:
            doc = fitz.open(stream=content, filetype="pdf")
            sections = []
            full_text = ""
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                page_text = page.get_text()
                
                if page_text.strip():
                    full_text += f"\n--- Page {page_num + 1} ---\n{page_text}"
                    
                    # Create sections
                    paragraphs = [p.strip() for p in page_text.split('\n\n') if p.strip()]
                    for para in paragraphs:
                        if len(para) > 50:
                            sections.append({
                                'page': page_num + 1,
                                'text': para,
                                'length': len(para),
                                'type': 'pdf_paragraph'
                            })
            
            doc.close()
            
            return {
                'full_text': full_text.strip(),
                'sections': sections,
                'page_count': len(doc),
                'format': 'PDF'
            }
            
        except Exception as e:
            return {'error': f'PDF processing error: {str(e)}'}
    
    def _process_docx_content(self, content: bytes) -> Dict[str, Any]:
        """Process DOCX content"""
        
        try:
            # Save content to temp file for python-docx
            import tempfile
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                tmp.write(content)
                tmp_path = tmp.name
            
            try:
                doc = Document(tmp_path)
                sections = []
                full_text = ""
                
                for i, paragraph in enumerate(doc.paragraphs):
                    text = paragraph.text.strip()
                    if text and len(text) > 20:
                        full_text += f"{text}\n\n"
                        sections.append({
                            'page': (i // 10) + 1,  # Approximate page
                            'text': text,
                            'length': len(text),
                            'type': 'docx_paragraph'
                        })
                
                # Process tables
                for table in doc.tables:
                    table_text = ""
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells])
                        table_text += row_text + "\n"
                    
                    if table_text.strip():
                        full_text += f"\n[TABLE]\n{table_text}\n"
                        sections.append({
                            'page': 1,
                            'text': table_text.strip(),
                            'length': len(table_text),
                            'type': 'docx_table'
                        })
                
                return {
                    'full_text': full_text.strip(),
                    'sections': sections,
                    'page_count': max(1, len(sections) // 10),
                    'format': 'DOCX'
                }
                
            finally:
                os.unlink(tmp_path)  # Clean up temp file
                
        except Exception as e:
            return {'error': f'DOCX processing error: {str(e)}'}
    
    def _process_email_content(self, content: bytes) -> Dict[str, Any]:
        """Process Email content"""
        
        try:
            # Parse email
            msg = email.message_from_bytes(content)
            
            sections = []
            full_text = ""
            
            # Extract headers
            headers = {
                'from': msg.get('From', ''),
                'to': msg.get('To', ''),
                'subject': msg.get('Subject', ''),
                'date': msg.get('Date', '')
            }
            
            header_text = f"From: {headers['from']}\nTo: {headers['to']}\nSubject: {headers['subject']}\nDate: {headers['date']}\n\n"
            full_text += header_text
            
            sections.append({
                'page': 1,
                'text': header_text.strip(),
                'length': len(header_text),
                'type': 'email_headers'
            })
            
            # Extract body
            body_text = ""
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        body_text += part.get_payload(decode=True).decode('utf-8', errors='ignore')
            else:
                body_text = msg.get_payload(decode=True).decode('utf-8', errors='ignore')
            
            if body_text.strip():
                full_text += body_text
                
                # Split email body into paragraphs
                paragraphs = [p.strip() for p in body_text.split('\n\n') if p.strip()]
                for para in paragraphs:
                    if len(para) > 30:
                        sections.append({
                            'page': 1,
                            'text': para,
                            'length': len(para),
                            'type': 'email_body'
                        })
            
            return {
                'full_text': full_text.strip(),
                'sections': sections,
                'page_count': 1,
                'format': 'EMAIL',
                'headers': headers
            }
            
        except Exception as e:
            return {'error': f'Email processing error: {str(e)}'}

# Export
__all__ = ['MultiFormatProcessor']